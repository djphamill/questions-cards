import csv
from docx import Document

def csv_to_docx(input_csv, output_docx):
    # Create a new Document
    doc = Document()

    # Open the CSV file
    with open(input_csv, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # Skip the header
        _ = next(reader)
        
        # Process each row in the CSV
        for row in reader:
            # Write each column of the row on a new line in the document
            name_paragraph = doc.add_paragraph('')
            name_paragraph.add_run(f'{row[0]} {row[1]}')

            doc.add_paragraph('')

            question_paragraph = doc.add_paragraph('')
            question = question_paragraph.add_run(f'{row[2]}')
            question.italic = True
            # Add a page break after each row
            doc.add_page_break()

    # Save the document
    doc.save(output_docx)
    print(f"Document saved as {output_docx}")

# Example usage
input_csv = 'input.csv'  # Replace with your CSV filename
output_docx = 'output.docx'  # The output document filename
csv_to_docx(input_csv, output_docx)

